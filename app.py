from flask import Flask, request, send_file, render_template
import zipfile
import io
import uuid
import os
from PyPDF2 import PdfMerger, PdfReader, PdfWriter
from pdf2docx import Converter
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from PIL import Image

app = Flask(__name__)

TEMP_DIR = "/tmp"
W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'


def temp_path(ext):
    return os.path.join(TEMP_DIR, str(uuid.uuid4()) + ext)


def is_run_bold(run_element):
    """Check if a run is bold - respects w:val='0' which means NOT bold."""
    rpr = run_element.find(f'{{{W_NS}}}rPr')
    if rpr is None:
        return False
    b_el = rpr.find(f'{{{W_NS}}}b')
    if b_el is None:
        return False
    # w:val="0" means explicitly NOT bold
    val = b_el.get(f'{{{W_NS}}}val')
    if val == '0' or val == 'false':
        return False
    return True


def get_run_font_size(run_element, default=11):
    """Get font size from run in points (sz is in half-points)."""
    rpr = run_element.find(f'{{{W_NS}}}rPr')
    if rpr is None:
        return default
    sz = rpr.find(f'{{{W_NS}}}sz')
    if sz is not None:
        val = sz.get(f'{{{W_NS}}}val')
        if val:
            return int(val) / 2
    return default


def get_para_default_size(para):
    """Get the most common font size in paragraph."""
    for child in para._element:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'r':
            sz = get_run_font_size(child)
            if sz != 11:
                return sz
    return 11


def build_rich_text(para):
    """Build rich text with per-run bold/italic/underline and hyperlinks."""
    result = ""
    default_size = get_para_default_size(para)

    for child in para._element:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'r':
            t_text = ''.join(t.text or '' for t in child.iter(f'{{{W_NS}}}t'))
            if not t_text:
                continue

            t_safe = t_text.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
            size = get_run_font_size(child, default_size)
            bold = is_run_bold(child)

            rpr = child.find(f'{{{W_NS}}}rPr')
            italic = rpr is not None and rpr.find(f'{{{W_NS}}}i') is not None
            underline = rpr is not None and rpr.find(f'{{{W_NS}}}u') is not None

            # Check italic is not explicitly turned off
            if italic:
                i_el = rpr.find(f'{{{W_NS}}}i')
                if i_el.get(f'{{{W_NS}}}val') in ('0', 'false'):
                    italic = False

            font_name = 'Helvetica-Bold' if bold else 'Helvetica'
            o, c = '', ''
            if bold:    o += '<b>'; c = '</b>' + c
            if italic:  o += '<i>'; c = '</i>' + c
            if underline: o += '<u>'; c = '</u>' + c

            result += f'<font name="{font_name}" size="{size}">{o}{t_safe}{c}</font>'

        elif tag == 'hyperlink':
            r_id = child.get(f'{{{R_NS}}}id')
            url = ''
            if r_id and r_id in para.part.rels:
                url = para.part.rels[r_id].target_ref

            hl_text = ''.join(t.text or '' for t in child.iter(f'{{{W_NS}}}t'))
            if not hl_text:
                continue

            hl_safe = hl_text.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
            if url:
                result += f'<font name="Helvetica" size="{default_size}" color="blue"><u><link href="{url}">{hl_safe}</link></u></font>'
            else:
                result += f'<font name="Helvetica" size="{default_size}" color="blue"><u>{hl_safe}</u></font>'

    return result


def get_full_text(para):
    """Get full paragraph text including hyperlinks."""
    text = ''
    for child in para._element:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag in ('r', 'hyperlink'):
            for t in child.iter(f'{{{W_NS}}}t'):
                text += (t.text or '')
    return text


def docx_to_pdf_proper(docx_path):
    doc = Document(docx_path)
    output = io.BytesIO()

    doc_template = SimpleDocTemplate(
        output, pagesize=A4,
        rightMargin=20*mm, leftMargin=20*mm,
        topMargin=20*mm, bottomMargin=20*mm
    )

    story = []

    for i, para in enumerate(doc.paragraphs):
        full_text = get_full_text(para)
        if not full_text.strip():
            story.append(Spacer(1, 3*mm))
            continue

        default_size = get_para_default_size(para)

        # Alignment
        alignment = TA_LEFT
        try:
            if para.alignment is not None:
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    alignment = TA_CENTER
                elif para.alignment == WD_ALIGN_PARAGRAPH.RIGHT:
                    alignment = TA_RIGHT
                elif para.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
                    alignment = TA_JUSTIFY
        except Exception:
            pass

        # Build rich text - bold is PER RUN not per paragraph
        rich_text = build_rich_text(para)
        if not rich_text:
            rich_text = full_text.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')

        # Paragraph style - NO bold at paragraph level
        ps = ParagraphStyle(
            name=f'p{i}',
            fontSize=default_size,
            leading=default_size * 1.35,
            alignment=alignment,
            fontName='Helvetica',
            spaceAfter=1.5*mm,
        )

        try:
            story.append(Paragraph(rich_text, ps))
        except Exception:
            plain = full_text.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
            try:
                story.append(Paragraph(plain, ps))
            except Exception:
                pass

    # Tables
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                cell_text = cell.text.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')
                row_data.append(Paragraph(cell_text, ParagraphStyle(
                    name='cell', fontSize=9, leading=12
                )))
            table_data.append(row_data)
        if table_data:
            col_count = max(len(r) for r in table_data)
            col_width = (A4[0] - 40*mm) / col_count
            t = Table(table_data, colWidths=[col_width]*col_count)
            t.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#f0f0ed')),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('FONTSIZE', (0,0), (-1,-1), 9),
                ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cccccc')),
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('PADDING', (0,0), (-1,-1), 4),
            ]))
            story.append(Spacer(1, 4*mm))
            story.append(t)
            story.append(Spacer(1, 4*mm))

    doc_template.build(story)
    output.seek(0)
    return output


@app.route("/", methods=["GET", "POST"])
def home():
    if request.method == "GET":
        return render_template("index.html")

    tool = request.form.get("tool")
    files = request.files.getlist("files")

    if not files or files[0].filename == "":
        return "No files uploaded"

    if tool == "merge_pdf":
        merger = PdfMerger()
        output = io.BytesIO()
        for f in files:
            merger.append(f)
        merger.write(output)
        merger.close()
        output.seek(0)
        return send_file(output, download_name="merged.pdf", as_attachment=True)

    elif tool == "compress_pdf":
        f = files[0]
        reader = PdfReader(f)
        writer = PdfWriter()
        for page in reader.pages:
            page.compress_content_streams()
            writer.add_page(page)
        if reader.metadata:
            writer.add_metadata(reader.metadata)
        output = io.BytesIO()
        writer.write(output)
        output.seek(0)
        original_name = f.filename or "file.pdf"
        download_name = original_name.replace(".pdf", "_compressed.pdf")
        return send_file(output, download_name=download_name, as_attachment=True)

    elif tool == "pdf_to_word":
        f = files[0]
        pdf_path = temp_path(".pdf")
        docx_path = temp_path(".docx")
        try:
            with open(pdf_path, "wb") as temp:
                temp.write(f.read())
            cv = Converter(pdf_path)
            cv.convert(docx_path)
            cv.close()
            with open(docx_path, "rb") as d:
                docx_buffer = io.BytesIO(d.read())
        finally:
            if os.path.exists(pdf_path): os.remove(pdf_path)
            if os.path.exists(docx_path): os.remove(docx_path)
        docx_buffer.seek(0)
        download_name = f.filename.replace(".pdf", ".docx")
        return send_file(docx_buffer, download_name=download_name, as_attachment=True)

    elif tool == "word_to_pdf":
        f = files[0]
        docx_path = temp_path(".docx")
        try:
            with open(docx_path, "wb") as temp:
                temp.write(f.read())
            pdf_buffer = docx_to_pdf_proper(docx_path)
        finally:
            if os.path.exists(docx_path): os.remove(docx_path)
        download_name = f.filename.replace(".docx", ".pdf")
        return send_file(pdf_buffer, download_name=download_name, as_attachment=True)

    elif tool == "extract_zip":
        uploaded_zip = zipfile.ZipFile(files[0])
        memory_zip = io.BytesIO()
        with zipfile.ZipFile(memory_zip, "w") as new_zip:
            for file in uploaded_zip.namelist():
                new_zip.writestr(file, uploaded_zip.read(file))
        memory_zip.seek(0)
        return send_file(memory_zip, download_name="extracted_files.zip", as_attachment=True)

    elif tool == "png_to_pdf":
        output = io.BytesIO()
        image_list = []
        for f in files:
            img = Image.open(f).convert("RGB")
            image_list.append(img)
        if len(image_list) == 1:
            image_list[0].save(output, format="PDF")
        else:
            image_list[0].save(output, format="PDF", save_all=True, append_images=image_list[1:])
        output.seek(0)
        first_name = files[0].filename or "image.png"
        download_name = first_name.rsplit(".", 1)[0] + ".pdf"
        return send_file(output, download_name=download_name, as_attachment=True)

    elif tool == "watermark_pdf":
        f = files[0]
        text = request.form.get("watermark_text", "CONFIDENTIAL")
        mode = request.form.get("watermark_mode", "watermark")
        reader = PdfReader(f)
        writer = PdfWriter()
        for page in reader.pages:
            page_width = float(page.mediabox.width)
            page_height = float(page.mediabox.height)
            overlay_buffer = io.BytesIO()
            c = canvas.Canvas(overlay_buffer, pagesize=(page_width, page_height))
            if mode == "watermark":
                c.setFont("Helvetica-Bold", 48)
                c.setFillColorRGB(0.7, 0.7, 0.7, alpha=0.35)
                c.saveState()
                c.translate(page_width / 2, page_height / 2)
                c.rotate(45)
                c.drawCentredString(0, 0, text)
                c.restoreState()
            else:
                c.setFont("Helvetica", 14)
                c.setFillColorRGB(0, 0, 0)
                c.drawString(40, 40, text)
            c.save()
            overlay_buffer.seek(0)
            overlay_reader = PdfReader(overlay_buffer)
            page.merge_page(overlay_reader.pages[0])
            writer.add_page(page)
        output = io.BytesIO()
        writer.write(output)
        output.seek(0)
        original_name = f.filename or "file.pdf"
        suffix = "_watermarked.pdf" if mode == "watermark" else "_edited.pdf"
        download_name = original_name.replace(".pdf", suffix)
        return send_file(output, download_name=download_name, as_attachment=True)

    return "Invalid Tool"


if __name__ == "__main__":
    app.run(debug=True)
